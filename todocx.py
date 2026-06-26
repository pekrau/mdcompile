"Compile Markdown files with footnotes, indexed terms and references to DOCX book."

import argparse
import datetime as dt
import os
import sys

# For debugging.
import icecream

icecream.install()

import docx
import docx.oxml
import docx.shared
import docx.styles.style
from docx.enum.style import WD_STYLE_TYPE

import constants
from text import Text
import utils


class Compiler:
    "Compile to DOCX format document."

    def __init__(
        self,
        filename,
        references=None,
        language=None,
        toc_level=None,
        page_break_level=None,
        text_number_level=None,
        no_comments=False,
        footnotes_location=None,
        paragraph_numbers=False,
    ):
        self.main = Text(filename)
        print(f"{len(self.main)} texts.")

        # Set the references directory.
        if references is None:
            try:
                references = os.environ["REFERENCES"]
            except KeyError:
                references = "./references"
        try:
            self.references = utils.ReferencesDir(references)
        except IOError:
            sys.exit(f"Error: no such reference directory '{references}'")

        # Set output parameters
        self.language = language or self.main.frontmatter.get(
            "language", constants.SV_SE
        )
        self.tx = utils.Tx(self.language)
        if toc_level is None:
            self.toc_level = self.main.frontmatter.get("toc-level", 1)
        else:
            self.toc_level = toc_level
        if page_break_level is None:
            self.page_break_level = self.main.frontmatter.get("page-break-level", 1)
        else:
            self.page_break_level = page_break_level
        if text_number_level is None:
            self.text_number_level = self.main.frontmatter.get("text-number-level", 1)
        else:
            self.text_number_level = text_number_level
        if no_comments:
            self.output_comments = False
        else:
            self.output_comments = self.main.frontmatter.get("output-comments", True)
        if paragraph_numbers or self.main.frontmatter.get("paragraph-numbers", False):
            self.paragraph_number = 0
        else:
            self.paragraph_number = None
        self.footnotes_location = footnotes_location or self.main.frontmatter.get(
            "footnotes-location", constants.FOOTNOTES_TEXT
        )

    def write(self, filename=None):
        "Convert the main text and its subtexts, if any, into DOCX."
        # Create and set up the DOCX document
        self.doc = docx.Document()

        # Set the default document-wide language.
        # From https://stackoverflow.com/questions/36967416/how-can-i-set-the-language-in-text-with-python-docx
        if self.language:
            styles_element = self.doc.styles.element
            rpr_default = styles_element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")[0]
            lang_default = rpr_default.xpath("w:lang")[0]
            lang_default.set(docx.oxml.shared.qn("w:val"), self.language)

        # Set to A4 page size.
        # XXX Make configurable.
        section = self.doc.sections[0]
        section.page_height = docx.shared.Mm(297)
        section.page_width = docx.shared.Mm(210)
        section.left_margin = docx.shared.Mm(25.4)
        section.right_margin = docx.shared.Mm(25.4)
        section.top_margin = docx.shared.Mm(25.4)
        section.bottom_margin = docx.shared.Mm(25.4)
        section.header_distance = docx.shared.Mm(12.7)
        section.footer_distance = docx.shared.Mm(12.7)

        # Create or modify styles to be used in the document.
        style = self.doc.styles.add_style("Title 0", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self.doc.styles["Title"]
        style.font.name = constants.DOCX_NORMAL_FONT
        style.font.size = docx.shared.Pt(constants.DOCX_FONT_SIZES[0])
        style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        for level in range(1, constants.MAX_LEVEL + 1):
            style = self.doc.styles.add_style(f"Title {level}", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = self.doc.styles[f"Heading {level}"]
            style.font.name = constants.DOCX_NORMAL_FONT
            style.font.size = docx.shared.Pt(constants.DOCX_FONT_SIZES[level])
            style.font.bold = True
            style.font.italic = False
            style.paragraph_format.space_before = docx.shared.Pt(
                5 * (constants.MAX_LEVEL + 1 - level)
            )
            style.paragraph_format.space_after = docx.shared.Pt(
                3 * (constants.MAX_LEVEL + 1 - level)
            )
            style.paragraph_format.line_spacing = 1
            style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        style = self.doc.styles["Normal"]
        style.font.name = constants.DOCX_NORMAL_FONT
        style.font.size = docx.shared.Pt(constants.DOCX_NORMAL_FONT_SIZE)
        style.paragraph_format.line_spacing = docx.shared.Pt(
            constants.DOCX_NORMAL_LINE_SPACING
        )

        # "Body Text": Table-of-contents (TOC) entries and index pages.
        style = self.doc.styles["Body Text"]
        style.paragraph_format.space_before = docx.shared.Pt(
            constants.DOCX_TOC_SPACE_BEFORE
        )
        style.paragraph_format.space_after = docx.shared.Pt(
            constants.DOCX_TOC_SPACE_AFTER
        )

        style = self.doc.styles["Quote"]
        style.paragraph_format.left_indent = docx.shared.Pt(constants.DOCX_QUOTE_INDENT)
        style.paragraph_format.right_indent = docx.shared.Pt(
            constants.DOCX_QUOTE_INDENT
        )

        style = self.doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = constants.DOCX_CODE_FONT
        style.font.size = docx.shared.Pt(constants.DOCX_CODE_FONT_SIZE)
        style.paragraph_format.line_spacing = docx.shared.Pt(
            constants.DOCX_CODE_LINE_SPACING
        )
        style.paragraph_format.left_indent = docx.shared.Pt(constants.DOCX_CODE_INDENT)

        style = self.doc.styles["Footer"]
        style.font.size = docx.shared.Pt(constants.DOCX_FOOTER_FONT_SIZE)

        # Define the footer contents.
        section.footer.paragraphs[0].text = (
            f"{self.tx('Created')}: {utils.isoformat()}\t{self.tx('Latest modification')}: {utils.isoformat(self.main.modified)}"
        )

        # Set Dublin core metadata.
        self.doc.core_properties.author = ", ".join(self.main.authors)
        self.doc.core_properties.created = dt.datetime.now()
        self.doc.core_properties.modified = self.main.modified
        if self.language:
            self.doc.core_properties.language = self.language

        # Display page number in the header.
        # https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
        paragraph = self.doc.sections[-1].header.paragraphs[0]
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        fldChar1 = docx.oxml.OxmlElement("w:fldChar")
        fldChar1.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
        instrText = docx.oxml.OxmlElement("w:instrText")
        instrText.set(docx.oxml.ns.qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = docx.oxml.OxmlElement("w:fldChar")
        fldChar2.set(docx.oxml.ns.qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Go through all elements in all texts to collect referenced and indexed.
        self.referenced = {}
        for text in self.main:
            for element in text.elements():
                if element["element"] != "reference":
                    continue
                if element["name"] not in self.referenced:
                    try:
                        self.referenced[element["name"]] = self.references[
                            element["name"]
                        ]
                    except KeyError:
                        sys.exit(f"Error: no such reference: '{element['name']}'")
        self.indexed = {}
        for text in self.main:
            for element in text.elements():
                if element["element"] != "indexed":
                    continue
                self.indexed.setdefault(element["canonical"], []).append(text)

        # Transfer footnotes to the appropriate texts, and number them.
        match self.footnotes_location:
            case constants.FOOTNOTES_TEXT:
                for text in self.main:
                    number = 0
                    for element in text.elements():
                        if element["element"] == "footnote_ref":
                            element["number"] = str(number := number + 1)
                            text.footnotes[element["label"]]["number"] = number
            case constants.FOOTNOTES_CHAPTER:
                for chapter in self.main.subtexts:
                    number = 0
                    for text in chapter:
                        for element in text.elements():
                            if element["element"] == "footnote_ref":
                                element["number"] = str(number := number + 1)
                                text.footnotes[element["label"]]["number"] = number
                        if text is not chapter:
                            labels = set(chapter.footnotes.keys()).intersection(
                                text.footnotes.keys()
                            )
                            if labels:
                                sys.exit(
                                    f"Error: footnote labels collision: {', '.join(labels)}"
                                )
                            chapter.footnotes.update(text.footnotes)
                            text.footnotes.clear()
            case constants.FOOTNOTES_BOOK:
                number = 0
                for text in self.main:
                    for element in text.elements():
                        if element["element"] == "footnote_ref":
                            element["number"] = str(number := number + 1)
                            text.footnotes[element["label"]]["number"] = number
                    if text is not self.main:
                        labels = set(self.main.footnotes.keys()).intersection(
                            text.footnotes.keys()
                        )
                        if labels:
                            sys.exit(
                                f"Error: footnote labels collision: {', '.join(labels)}"
                            )
                        self.main.footnotes.update(text.footnotes)
                        text.footnotes.clear()

        # 0: not in footnote; -1: footnote started; >= 1: footnote number to start
        self.footnote_def_flag = 0
        print(f"Footnotes at end of {self.footnotes_location}.")

        # Output title page.
        paragraph = self.doc.add_paragraph(style="Title 0")
        run = paragraph.add_run(self.main.title)

        if self.main.subtitle:
            paragraph = self.doc.add_paragraph(style="Title 1")
            paragraph.add_run(self.main.subtitle)

        # Split authors into runs to allow line break between them.
        paragraph = self.doc.add_paragraph(style="Title 2")
        for author in self.main.authors:
            paragraph.add_run(author)
            if author != self.main.authors[-1]:
                paragraph.add_run(", ")

        # Title-page text; synopsis, or similar.
        Renderer(self, self.main)

        # Write table of contents (TOC) page(s).
        # The DOCX format does not allow determining the page numbers before printing.
        if self.toc_level:
            self.doc.add_page_break()
            self.write_heading(self.tx("Contents"), 1)
            for text in list(self.main)[1:]:  # Skip the main file; title page.
                if text.level > self.toc_level:
                    continue
                paragraph = self.doc.add_paragraph(style="Body Text")
                paragraph.paragraph_format.left_indent = docx.shared.Pt(
                    constants.DOCX_TOC_INDENT * text.level
                )
                paragraph.paragraph_format.first_line_indent = -docx.shared.Pt(
                    constants.DOCX_TOC_INDENT
                )
                paragraph.add_run(self.numbered_title(text))

            # Output entries for book footnotes, references and indexed, if any such.
            if (
                self.footnotes_location == constants.FOOTNOTES_BOOK
                and self.main.footnotes
            ):
                self.doc.add_paragraph(self.tx("Footnotes"), style="Body Text")
            if self.referenced:
                self.doc.add_paragraph(self.tx("References"), style="Body Text")
            if self.indexed:
                self.doc.add_paragraph(self.tx("Index"), style="Body Text")

        # First-level subtexts are chapters.
        for text in self.main.subtexts:
            self.write_text(text)
            if (
                self.footnotes_location == constants.FOOTNOTES_CHAPTER
                and text.footnotes
            ):
                self.doc.add_page_break()
                self.write_heading(self.tx("Footnotes"), 3)
                self.write_footnotes(text)

        if self.footnotes_location == constants.FOOTNOTES_BOOK:
            self.doc.add_page_break()
            self.write_heading(self.tx("Footnotes"), 1)
            self.write_footnotes(self.main)

        self.write_referenced()
        print(f"{len(self.referenced)} references used.")

        self.write_indexed()
        print(f"{len(self.indexed)} terms indexed.")

        if self.paragraph_number:
            print(f"{self.paragraph_number} paragraphs.")

        filename = filename or self.main.filename.with_suffix(".docx")
        self.doc.save(filename)

    def write_text(self, text):
        "Output the contents of the text instance."
        if text.level <= self.page_break_level:
            self.doc.add_page_break()
        self.write_heading(self.numbered_title(text), text.level)
        if text.subtitle:
            self.write_heading(text.subtitle, text.level + 1)
        self.current_text = text

        Renderer(self, text)

        if self.footnotes_location == constants.FOOTNOTES_TEXT and text.footnotes:
            self.write_heading(self.tx("Footnotes"), text.level + 2)
            self.write_footnotes(text)

        for subtext in text.subtexts:
            self.write_text(subtext)

    def write_footnotes(self, text):
        "Write out the footnotes for the text."
        for footnote in sorted(text.footnotes.values(), key=lambda f: f["number"]):
            self.footnote_def_flag = footnote["number"]
            for child in footnote["children"]:
                Renderer(self, ast=child)
            self.footnote_def_flag = 0

    def write_heading(self, heading, level):
        if level <= constants.MAX_LEVEL:
            paragraph = self.doc.add_paragraph(style=f"Title {level}")
            paragraph.add_run(heading)
        else:
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(heading)
            run.font.italic = True

    def numbered_title(self, text, force=False):
        "Return the title with a number prefix."
        if text.level <= self.text_number_level or force:
            return f"{'.'.join([str(i) for i in text.ordinal])}. {text.title}"
        else:
            return text.title

    def write_referenced(self):
        "Write referenced pages, if any such items."
        if not self.referenced:
            return
        self.doc.add_page_break()
        self.write_heading(self.tx("References"), 1)
        for name, reference in sorted(self.referenced.items()):
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.left_indent = docx.shared.Pt(
                constants.DOCX_REFERENCE_INDENT
            )
            paragraph.paragraph_format.first_line_indent = -docx.shared.Pt(
                constants.DOCX_REFERENCE_INDENT
            )
            run = paragraph.add_run(reference["reference"])
            run.font.bold = True
            paragraph.add_run("  ")
            self.write_reference_authors(paragraph, reference)
            try:
                method = getattr(self, f"write_reference_{reference['type']}")
            except AttributeError:
                sys.exit(f"Error: unknown reference type {reference['type']}")
            else:
                method(paragraph, reference)
            paragraph.add_run(".")
            self.write_reference_external_links(paragraph, reference)

    def write_reference_authors(self, paragraph, reference):
        count = len(reference["authors"])
        for pos, author in enumerate(reference["authors"]):
            if pos > 0:
                if pos == count - 1:
                    paragraph.add_run(" & ")
                else:
                    paragraph.add_run(", ")
            paragraph.add_run(utils.short_person_name(author))

    def write_reference_article(self, paragraph, reference):
        paragraph.add_run(" ")
        paragraph.add_run(reference["title"])
        try:
            run = paragraph.add_run(f", {reference['journal']}")
            run.font.italic = True
        except KeyError:
            pass
        try:
            paragraph.add_run(f" {reference['volume']}")
        except KeyError:
            pass
        else:
            try:
                paragraph.add_run(f" ({reference['issue']})")
            except KeyError:
                pass
        try:
            paragraph.add_run(f" {reference['pages'].replace('--', '-')}")
        except KeyError:
            pass

    def write_reference_book(self, paragraph, reference):
        paragraph.add_run(" ")
        run = paragraph.add_run(reference["title"])
        run.font.italic = True
        try:
            paragraph.add_run(f" {reference['publisher']}")
            paragraph.add_run(f", {reference['edition_published']}")
        except KeyError:
            pass

    def write_reference_link(self, paragraph, reference):
        paragraph.add_run(" ")
        run = paragraph.add_run(reference["title"])
        run.font.italic = True
        if url := reference.get("url"):
            paragraph.add_run(" ")
            self.add_hyperlink(paragraph, url, "")
            try:
                paragraph.add_run(f" Accessed {reference['accessed']}.")
            except KeyError:
                pass

    def write_reference_external_links(self, paragraph, reference):
        any_item = False
        if url := reference.get("url"):
            self.add_hyperlink(paragraph, url, url)
            any_item = True
        for key, (label, template) in constants.REFS_LINKS.items():
            try:
                value = reference[key]
                if any_item:
                    paragraph.add_run(", ")
                else:
                    paragraph.add_run(" ")
                self.add_hyperlink(
                    paragraph, template.format(value=value), f"{label}:{value}"
                )
                any_item = True
            except KeyError:
                pass

    def write_indexed(self):
        "Write indexed terms pages, if any such items."
        if not self.indexed:
            return
        self.doc.add_page_break()
        self.write_heading(self.tx("Index"), 1)
        items = sorted(self.indexed.items(), key=lambda i: i[0].casefold())
        for canonical, texts in items:
            paragraph = self.doc.add_paragraph(canonical, style="Body Text")
            paragraph.paragraph_format.keep_with_next = True
            for text in texts:
                paragraph = self.doc.add_paragraph(self.numbered_title(text, force=True), style="Body Text")
                paragraph.paragraph_format.left_indent = docx.shared.Pt(
                    constants.DOCX_INDEXED_INDENT
                )
                if text is not texts[:-1]:
                    paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_after = docx.shared.Pt(
                constants.DOCX_INDEXED_SPACE_AFTER
            )

    # https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
    def add_hyperlink(self, paragraph, url, text, color="2222FF", underline=True):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        # Get access to the document.xml.rels file and gets a new relation id value.
        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )

        # Create the w:hyperlink tag and add needed values.
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(
            docx.oxml.shared.qn("r:id"),
            r_id,
        )

        # Create a w:r element.
        new_run = docx.oxml.shared.OxmlElement("w:r")

        # Create a new w:rPr element.
        rPr = docx.oxml.shared.OxmlElement("w:rPr")

        # Add color if it is given.
        if not color is None:
            c = docx.oxml.shared.OxmlElement("w:color")
            c.set(docx.oxml.shared.qn("w:val"), color)
            rPr.append(c)

        # Remove underlining if it is requested.
        # XXX Does not seem to work? /Per Kraulis
        if not underline:
            u = docx.oxml.shared.OxmlElement("w:u")
            u.set(docx.oxml.shared.qn("w:val"), "none")
            rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element.
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink


class Renderer:
    "Render the Markdown text AST."

    def __init__(self, compiler, text=None, ast=None):
        assert text is not None or ast is not None
        self.compiler = compiler
        self.doc = compiler.doc
        self.text = text
        self.list_stack = []
        self.style_stack = ["Normal"]
        self.bold = False
        self.italic = False
        self.subscript = False
        self.superscript = False
        self(ast or text.ast)

    def __call__(self, ast):
        "Render the Markdown text AST."
        try:
            method = getattr(self, ast["element"])
        except AttributeError:
            print("Could not handle ast", ast)
        else:
            method(ast)

    def document(self, ast):
        self.prev_blank_line = False
        for child in ast["children"]:
            self(child)

    def heading(self, ast):
        # XXX Limited implementation; this just handles one child of raw text.
        level = ast["level"]
        if self.text:
            level += self.text.level
        self.compiler.write_heading(ast["children"][0]["children"], level)

    def paragraph(self, ast):
        self.current_paragraph = self.doc.add_paragraph()

        # Either starting footnote definition, or within it.
        if self.compiler.footnote_def_flag != 0:
            self.current_paragraph.paragraph_format.left_indent = docx.shared.Pt(
                constants.DOCX_FOOTNOTE_INDENT
            )
            # Starting footnote definition.
            if self.compiler.footnote_def_flag > 0:
                self.current_paragraph.paragraph_format.first_line_indent = (
                    -docx.shared.Pt(constants.DOCX_FOOTNOTE_INDENT)
                )
                run = self.current_paragraph.add_run(
                    f"{self.compiler.footnote_def_flag}."
                )
                run.font.bold = True
                self.current_paragraph.add_run(" ")
                # Signal for being within footnote definition.
                self.compiler.footnote_def_flag = -1

        if self.list_stack:
            data = self.list_stack[-1]
            levels = min(3, data["levels"])  # Max list levels in predef list styles.
            if data["first_paragraph"]:
                if data["ordered"]:
                    if levels == 1:
                        style = self.doc.styles["List Number"]
                    else:
                        style = self.doc.styles[f"List Number {levels}"]
                else:
                    if levels == 1:
                        style = self.doc.styles["List Bullet"]
                    else:
                        style = self.doc.styles[f"List Bullet {levels}"]
            else:
                if levels == 1:
                    style = self.doc.styles["List Continue"]
                else:
                    style = self.doc.styles[f"List Continue {levels}"]
            data["first_paragraph"] = False
            self.current_paragraph.style = style
        else:
            self.current_paragraph.style = self.style_stack[-1]

        if self.compiler.paragraph_number is not None:
            self.compiler.paragraph_number += 1
            run = self.current_paragraph.add_run(f"{self.compiler.paragraph_number}. ")
            run.style = self.doc.styles["Intense Quote Char"]

        for child in ast["children"]:
            self(child)

    def raw_text(self, ast):
        line = ast["children"]
        line = line.rstrip("\n")
        run = self.current_paragraph.add_run(line)
        if self.bold:
            run.font.bold = True
        if self.italic:
            run.font.italic = True
        if self.subscript:
            run.font.subscript = True
        if self.superscript:
            run.font.superscript = True

    def blank_line(self, ast):
        pass

    def quote(self, ast):
        self.style_stack.append("Quote")
        for child in ast["children"]:
            self(child)
        self.style_stack.pop()

    def code_span(self, ast):
        run = self.current_paragraph.add_run(ast["children"])
        run.style = self.doc.styles["Macro Text Char"]

    def code_block(self, ast):
        self.current_paragraph = self.doc.add_paragraph(style="Code")
        self.style_stack.append("Code")
        for child in ast["children"]:
            self(child)
        self.style_stack.pop()

    def fenced_code(self, ast):
        self.current_paragraph = self.doc.add_paragraph(style="Code")
        self.style_stack.append("Code")
        for child in ast["children"]:
            self(child)
        self.style_stack.pop()

    def image(self, ast):
        try:
            # Fetch image from the web.
            if urllib.parse.urlparse(ast["dest"]).scheme:
                response = requests.get(ast["dest"])
                if response.status_code != HTTP.OK:
                    raise ValueError(f"Could not fetch image '{ast['dest']}'")
                if response.headers["Content-Type"] not in (
                    constants.PNG_MIMETYPE,
                    constants.JPEG_MIMETYPE,
                ):
                    raise ValueError(
                        f"Cannot handle image '{ast['dest']}' with content type '{response.headers['Content-Type']}'"
                    )
                self.add_image(response.content, ast, 1.0)

            # Use image from the image library.
            elif ast["dest"] in get_imgs():
                img = get_imgs()[ast["dest"]]
                scale_factor = img["docx"]["scale_factor"]

                if img["content_type"] in (
                    constants.SVG_MIMETYPE,
                    constants.JSON_MIMETYPE,
                ):
                    # SVG image.
                    if img["content_type"] == constants.SVG_MIMETYPE:
                        # SVG in image library has already been checked for validity.
                        root = minixml.parse_content(img["data"])

                    # Vega-Lite plot.
                    else:
                        # JSON in image library has already been checked for validity.
                        vl_spec = json.loads(img["data"])
                        root = minixml.parse_content(
                            vl_convert.vegalite_to_svg(vl_spec)
                        )

                    # Set viewbox so that scaling behaves.
                    root["viewBox"] = f"0 0 {root['width']} {root['height']}"

                    # Scale width and height in SVG element.
                    rendering_factor = img["docx"]["png_rendering_factor"]
                    root["width"] = rendering_factor * float(root["width"])
                    root["height"] = rendering_factor * float(root["height"])

                    self.add_image(
                        vl_convert.svg_to_png(repr(root)),
                        ast,
                        scale_factor / rendering_factor,
                    )

                # JPEG or PNG.
                elif img["content_type"] in (
                    constants.PNG_MIMETYPE,
                    constants.JPEG_MIMETYPE,
                ):
                    self.add_image(
                        base64.standard_b64decode(img["data"]), ast, scale_factor
                    )
                else:
                    raise ValueError(f"Cannot handle image {img['content_type']}")
            else:
                raise ValueError(f"No such image '{ast['dest']}'")

        except ValueError as error:
            self.current_paragraph = self.doc.add_paragraph(style="Code")
            self.current_paragraph.add_run(str(error))

    def add_image(self, image_data, ast, factor):
        image = io.BytesIO(image_data)
        width, height = PIL.Image.open(image).size
        width = docx.shared.Pt(factor * width)
        height = docx.shared.Pt(factor * height)
        paragraph = self.doc.add_paragraph()
        # This is a kludge; seems required to avoid an obscure 'docx' bug?
        paragraph.paragraph_format.line_spacing = 1
        paragraph.add_run().add_picture(image, width=width, height=height)
        if ast["children"]:
            paragraph.paragraph_format.keep_with_next = True
            self.current_paragraph = self.doc.add_paragraph(style="Normal")
            for child in ast["children"]:
                self(child)

    def emphasis(self, ast):
        self.italic = True
        for child in ast["children"]:
            self(child)
        self.italic = False

    def strong_emphasis(self, ast):
        self.bold = True
        for child in ast["children"]:
            self(child)
        self.bold = False

    def subscript(self, ast):
        self.subscript = True
        for child in ast["children"]:
            self(child)
        self.subscript = False

    def superscript(self, ast):
        self.superscript = True
        for child in ast["children"]:
            self(child)
        self.superscript = False

    def emdash(self, ast):
        self.current_paragraph.add_run(constants.EM_DASH)

    def line_break(self, ast):
        if ast.get("soft"):
            self.current_paragraph.add_run(" ")
        else:
            self.current_paragraph.add_run("\n")

    def thematic_break(self, ast):
        paragraph = self.doc.add_paragraph(constants.EM_DASH * 20)
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def link(self, ast):
        # This handles only raw text within a link; can't handle multiple children.
        raw_text = []
        for child in ast["children"]:
            if child["element"] == "raw_text":
                raw_text.append(child["children"])
        self.compiler.add_hyperlink(
            self.current_paragraph, ast["dest"], "".join(raw_text)
        )

    def list(self, ast):
        data = dict(
            ordered=ast["ordered"],
            bullet=ast["bullet"],  # Currently not used.
            start=ast["start"],  # Currently not used.
            tight=ast["tight"],  # Currently not used.
            count=0,  # Currently not used.
            levels=len(self.list_stack) + 1,
        )
        self.list_stack.append(data)
        for child in ast["children"]:
            self(child)
        self.list_stack.pop()

    def list_item(self, ast):
        data = self.list_stack[-1]
        data["count"] += 1  # Currently not used.
        data["first_paragraph"] = True
        for child in ast["children"]:
            self(child)

    def indexed(self, ast):
        run = self.current_paragraph.add_run(ast["term"])
        run.font.underline = True

    def footnote_ref(self, ast):
        run = self.current_paragraph.add_run(f" {ast['number']}")
        run.font.superscript = True
        run.font.bold = True

    def footnote_def(self, ast):
        "The footnote definition in the element stream is not used; ignore."
        pass

    def reference(self, ast):
        run = self.current_paragraph.add_run(ast["name"])
        run.font.bold = True

    def comment(self, ast):
        if self.output_comments:
            run = self.current_paragraph.add_run(ast["comment"])
            run.font.bold = True
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    def link_ref_def(self, ast):
        pass


if __name__ == "__main__":
    args = utils.get_args("todocx")
    compiler = Compiler(**vars(args))
    compiler.write()
