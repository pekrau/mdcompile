"Compile Markdown file(s) with footnotes, indexed terms and references to DOCX."

import datetime as dt

import docx
import docx.oxml
import docx.shared
import docx.styles.style
from docx.enum.style import WD_STYLE_TYPE

import constants
from compiler import Compiler
from text import Text
import utils


class DocxCompiler(Compiler):
    "Compile to DOCX format."

    def write(self, filename=None):
        "Convert the main text and its subtexts, if any, into DOCX."
        # Create and set up the DOCX document
        self.doc = docx.Document()

        # Set the default document-wide language.
        # From https://stackoverflow.com/questions/36967416/how-can-i-set-the-language-in-text-with-python-docx
        styles_element = self.doc.styles.element
        rpr_default = styles_element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")[0]
        lang_default = rpr_default.xpath("w:lang")[0]
        lang_default.set(docx.oxml.shared.qn("w:val"), self.main.language)

        # Set to A4 page size. XXX Allow alternatives.
        section = self.doc.sections[0]
        section.page_height = docx.shared.Mm(297)
        section.page_width = docx.shared.Mm(210)
        section.left_margin = docx.shared.Mm(25.4)
        section.right_margin = docx.shared.Mm(25.4)
        section.top_margin = docx.shared.Mm(25.4)
        section.bottom_margin = docx.shared.Mm(25.4)
        section.header_distance = docx.shared.Mm(12.7)
        section.footer_distance = docx.shared.Mm(12.7)

        # Create or modify styles to be used in the DOCX document.
        style = self.doc.styles.add_style("Title 0", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self.doc.styles["Title"]
        style.font.name = constants.DOCX_NORMAL_FONT
        style.font.size = docx.shared.Pt(constants.DOCX_FONT_SIZES[0])
        style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        for level in range(1, constants.MAX_LEVEL + 1):
            style = self.doc.styles.add_style(f"Title {level}", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = self.doc.styles[f"Heading {level}"]
            style.font.name = constants.DOCX_NORMAL_FONT
            style.font.size = docx.shared.Pt(constants.DOCX_FONT_SIZES[level])
            style.font.bold = True
            style.font.italic = False
            style.paragraph_format.space_before = docx.shared.Pt(
                5 * (constants.MAX_LEVEL + 1 - level)
            )
            style.paragraph_format.space_after = docx.shared.Pt(
                3 * (constants.MAX_LEVEL + 1 - level)
            )
            style.paragraph_format.line_spacing = 1
            style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        style = self.doc.styles["Normal"]
        style.font.name = constants.DOCX_NORMAL_FONT
        style.font.size = docx.shared.Pt(constants.DOCX_NORMAL_FONT_SIZE)
        style.paragraph_format.line_spacing = docx.shared.Pt(
            constants.DOCX_NORMAL_LINE_SPACING
        )

        # "Body Text": Table-of-contents (TOC) entries and index pages.
        style = self.doc.styles["Body Text"]
        style.paragraph_format.space_before = docx.shared.Pt(
            constants.DOCX_TOC_SPACE_BEFORE
        )
        style.paragraph_format.space_after = docx.shared.Pt(
            constants.DOCX_TOC_SPACE_AFTER
        )

        style = self.doc.styles["Quote"]
        style.paragraph_format.left_indent = docx.shared.Pt(constants.DOCX_QUOTE_INDENT)
        style.paragraph_format.right_indent = docx.shared.Pt(
            constants.DOCX_QUOTE_INDENT
        )
        style.font.size = docx.shared.Pt(constants.DOCX_QUOTE_FONT_SIZE)

        style = self.doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = constants.DOCX_CODE_FONT
        style.font.size = docx.shared.Pt(constants.DOCX_CODE_FONT_SIZE)
        style.paragraph_format.line_spacing = docx.shared.Pt(
            constants.DOCX_CODE_LINE_SPACING
        )
        style.paragraph_format.left_indent = docx.shared.Pt(constants.DOCX_CODE_INDENT)

        # Set Dublin core metadata.
        self.doc.core_properties.author = ", ".join(self.main.authors)
        self.doc.core_properties.created = dt.datetime.now()
        self.doc.core_properties.modified = self.main.modified
        self.doc.core_properties.language = self.main.language

        # Display page number in the DOCX header.
        # https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
        paragraph = self.doc.sections[-1].header.paragraphs[0]
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        fldChar1 = docx.oxml.OxmlElement("w:fldChar")
        fldChar1.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
        instrText = docx.oxml.OxmlElement("w:instrText")
        instrText.set(docx.oxml.ns.qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = docx.oxml.OxmlElement("w:fldChar")
        fldChar2.set(docx.oxml.ns.qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Write title page: title, subtitle, initial note, authors, dates.
        paragraph = self.doc.add_paragraph(style="Title 0")
        paragraph.add_run(self.main.title)
        if self.main.subtitle:
            paragraph = self.doc.add_paragraph(style="Title 1")
            paragraph.add_run(self.main.subtitle)

        # Split authors into runs to allow line break between them.
        paragraph = self.doc.add_paragraph(style="Title 2")
        for author in self.main.authors:
            paragraph.add_run(author)
            if author != self.main.authors[-1]:
                paragraph.add_run(", ")

        self.text_render(self.main)

        paragraph = self.doc.add_paragraph(style="Normal")
        paragraph.add_run(
            f'{self.tx("Created")}: {utils.isoformat()}\n'
            f'{self.tx("Latest modification")}: {utils.isoformat(self.main.modified)}'
        ).font.italic = True

        # Write table of contents (TOC) page(s).
        # The DOCX format does not allow determining the page numbers before printing.
        if self.toc_level and self.main.subtexts:
            self.write_page_break()
            self.write_heading(self.tx("Contents"), 1)
            for text in list(self.main)[1:]:  # Skip the main file (title page).
                if text.level <= self.toc_level:
                    paragraph = self.doc.add_paragraph(style="Body Text")
                    paragraph.paragraph_format.left_indent = docx.shared.Pt(
                        constants.DOCX_TOC_INDENT * text.level
                    )
                    paragraph.paragraph_format.first_line_indent = -docx.shared.Pt(
                        constants.DOCX_TOC_INDENT
                    )
                    if text.level <= self.text_number_level:
                        title = text.ordinal_title
                    else:
                        title = text.title
                    paragraph.add_run(title)

            # Write entries for book footnotes, references and indexed, if any such.
            if (
                self.footnotes_location == constants.FOOTNOTES_BOOK
                and self.main.footnotes
            ):
                self.doc.add_paragraph(self.tx("Footnotes"), style="Body Text")
            if self.referenced:
                self.doc.add_paragraph(self.tx("References"), style="Body Text")
            if self.indexed:
                self.doc.add_paragraph(self.tx("Index"), style="Body Text")

        # First-level subtexts are chapters.
        for text in self.main.subtexts:
            self.write_text(text)
            if self.footnotes_location == constants.FOOTNOTES_CHAPTER:
                if text.footnotes:
                    self.write_page_break()
                    self.write_heading(self.tx("Footnotes"), 3)
                    self.write_footnotes(text)

        if self.footnotes_location == constants.FOOTNOTES_BOOK:
            self.write_page_break()
            self.write_heading(self.tx("Footnotes"), 1)
            self.write_footnotes(self.main)

        self.write_referenced()
        self.write_indexed()

        filename = filename or self.main.filename.with_suffix(".docx")
        self.doc.save(filename)

    def write_text(self, text):
        "Write the contents of the text instance."
        if text.level <= self.page_break_level:
            self.write_page_break()
        if text.level <= self.text_number_level:
            title = text.ordinal_title
        else:
            title = text.title
        self.write_heading(title, text.level)
        if text.subtitle:
            self.write_heading(text.subtitle, text.level + 1)

        self.text_render(text)

        if self.footnotes_location == constants.FOOTNOTES_TEXT and text.footnotes:
            self.write_heading(self.tx("Footnotes"), text.level + 2)
            self.write_footnotes(text)

        for subtext in text.subtexts:
            self.write_text(subtext)

    def write_heading(self, heading, level):
        if level <= constants.MAX_LEVEL:
            paragraph = self.doc.add_paragraph(style=f"Title {level}")
            paragraph.add_run(heading)
        else:
            paragraph = self.doc.add_paragraph()
            paragraph.add_run(heading).font.italic = True

    def write_page_break(self):
        self.doc.add_page_break()

    def write_footnotes(self, text):
        "Write out the footnotes for the text."
        for footnote in sorted(text.footnotes.values(), key=lambda f: f["number"]):
            self.footnote_def_flag = footnote["number"]
            for child in footnote["children"]:
                self.render(child)
            self.footnote_def_flag = 0

    def write_referenced(self):
        "Write referenced pages, if any."
        if not self.referenced:
            return
        self.write_page_break()
        self.write_heading(self.tx("References"), 1)
        for name, reference in sorted(self.referenced.items()):
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.left_indent = docx.shared.Pt(
                constants.DOCX_REFERENCE_INDENT
            )
            paragraph.paragraph_format.first_line_indent = -docx.shared.Pt(
                constants.DOCX_REFERENCE_INDENT
            )
            paragraph.add_run(name).font.bold = True
            paragraph.add_run("  ")
            self.write_reference_authors(paragraph, reference)
            try:
                method = getattr(self, f"write_reference_{reference['type']}")
            except AttributeError:
                raise ValueError(f"unknown reference type {reference['type']}")
            else:
                method(paragraph, reference)
            paragraph.add_run(".")
            self.write_reference_external_links(paragraph, reference)

    def write_reference_authors(self, paragraph, reference):
        count = len(reference["authors"])
        for pos, author in enumerate(reference["authors"]):
            if pos > 0:
                if pos == count - 1:
                    paragraph.add_run(" & ")
                else:
                    paragraph.add_run(", ")
            paragraph.add_run(author)
        paragraph.add_run(": ")

    def write_reference_article(self, paragraph, reference):
        paragraph.add_run(reference["title"])
        paragraph.add_run(",")
        if journal := reference.get("journal"):
            paragraph.add_run(" ")
            paragraph.add_run(journal).font.italic = True
        if volume := reference.get("volume"):
            paragraph.add_run(" ")
            paragraph.add_run(volume).font.bold = True
        if issue := reference.get("issue"):
            paragraph.add_run(f" ({issue})")
        if pages := reference.get("pages"):
            paragraph.add_run(" ")
            paragraph.add_run(pages.replace("--", "-"))

    def write_reference_book(self, paragraph, reference):
        paragraph.add_run(reference["title"]).font.italic = True
        if edition := reference.get("edition"):
            paragraph.add_run(",")
            if publisher := edition.get("publisher"):
                paragraph.add_run(" ")
                paragraph.add_run(publisher)
            if published := edition.get("published"):
                paragraph.add_run(" ")
                paragraph.add_run(published)

    def write_reference_link(self, paragraph, reference):
        "Write reference which is a URL to a website."
        paragraph.add_run(reference["title"]).font.italic = True
        if url := reference.get("url"):
            paragraph.add_run(", ")
            self.add_hyperlink(paragraph, url, url)
            if accessed := reference.get("accessed"):
                paragraph.add_run(f" ({self.tx('accessed')} {accessed})")

    def write_reference_external_links(self, paragraph, reference):
        "Write external links; doi, pmid, isbn, ..."
        if url := reference.get("url"):
            paragraph.add_run(" ")
            self.add_hyperlink(paragraph, url, url)
        for key, (label, template) in constants.REFS_LINKS.items():
            try:
                value = reference[key]
                url = template.format(value=value)
                paragraph.add_run(" ")
                self.add_hyperlink(paragraph, url, f"{label}:{value}")
            except KeyError:
                pass

    def write_indexed(self):
        "Write indexed terms pages, if any such items."
        if not self.indexed:
            return
        self.write_page_break()
        self.write_heading(self.tx("Index"), 1)
        items = sorted(self.indexed.items(), key=lambda i: i[0].casefold())
        for canonical, texts in items:
            paragraph = self.doc.add_paragraph(canonical, style="Body Text")
            paragraph.paragraph_format.keep_with_next = True
            for text in texts:
                paragraph = self.doc.add_paragraph(text.ordinal_title,style="Body Text")
                paragraph.paragraph_format.left_indent = docx.shared.Pt(
                    constants.DOCX_INDEXED_INDENT
                )
                if text is not texts[:-1]:
                    paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_after = docx.shared.Pt(
                constants.DOCX_INDEXED_SPACE_AFTER
            )

    # https://github.com/python-openxml/python-docx/issues/610
    def add_hyperlink(self, paragraph, url, text):
        run = paragraph.add_run(text)
        # Gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)
        hyperlink.append(run._r)
        paragraph._p.insert(len(paragraph.runs) + 2, hyperlink)
        run.font.color.rgb = docx.shared.RGBColor(65, 105, 225)
        run.font.underline = True

    def add_image(self, image_data, ast, factor):
        image = io.BytesIO(image_data)
        width, height = PIL.Image.open(image).size
        width = docx.shared.Pt(factor * width)
        height = docx.shared.Pt(factor * height)
        paragraph = self.doc.add_paragraph()
        # This is a kludge; seems required to avoid an obscure 'docx' bug?
        paragraph.paragraph_format.line_spacing = 1
        paragraph.add_run().add_picture(image, width=width, height=height)
        if ast["children"]:
            paragraph.paragraph_format.keep_with_next = True
            self.current_paragraph = self.doc.add_paragraph(style="Normal")
            for child in ast["children"]:
                self.render(child)

    def render_document(self, ast):
        self.prev_blank_line = False
        for child in ast["children"]:
            self.render(child)

    def render_heading(self, ast):
        # XXX Limited implementation; this just handles one child of raw text.
        level = ast["level"]
        if self.current_text:
            level += self.current_text.level
        self.write_heading(ast["children"][0]["children"], level)

    def render_paragraph(self, ast):
        self.current_paragraph = self.doc.add_paragraph()

        # Either starting footnote definition, or within it.
        if self.footnote_def_flag != 0:
            self.current_paragraph.paragraph_format.left_indent = docx.shared.Pt(
                constants.DOCX_FOOTNOTE_INDENT
            )
            # Starting footnote definition.
            if self.footnote_def_flag > 0:
                self.current_paragraph.paragraph_format.first_line_indent = (
                    -docx.shared.Pt(constants.DOCX_FOOTNOTE_INDENT)
                )
                self.current_paragraph.add_run(
                    f"{self.footnote_def_flag}."
                ).font.bold = True
                self.current_paragraph.add_run(" ")
                # Signal for being within footnote definition.
                self.footnote_def_flag = -1

        if self.list_stack:
            data = self.list_stack[-1]
            levels = min(3, data["levels"])  # Max list levels in predef list styles.
            if data["first_paragraph"]:
                if data["ordered"]:
                    if levels == 1:
                        style = self.doc.styles["List Number"]
                    else:
                        style = self.doc.styles[f"List Number {levels}"]
                else:
                    if levels == 1:
                        style = self.doc.styles["List Bullet"]
                    else:
                        style = self.doc.styles[f"List Bullet {levels}"]
            else:
                if levels == 1:
                    style = self.doc.styles["List Continue"]
                else:
                    style = self.doc.styles[f"List Continue {levels}"]
            data["first_paragraph"] = False
            self.current_paragraph.style = style
        else:
            self.current_paragraph.style = self.style_stack[-1]

        if self.paragraph_number is not None:
            self.paragraph_number += 1
            self.current_paragraph.add_run(f"{self.paragraph_number}. ").style = (
                self.doc.styles["Intense Quote Char"]
            )

        for child in ast["children"]:
            self.render(child)

    def render_raw_text(self, ast):
        line = ast["children"]
        line = line.rstrip("\n")
        run = self.current_paragraph.add_run(line)
        if self.bold:
            run.font.bold = True
        if self.italic:
            run.font.italic = True
        if self.subscript:
            run.font.subscript = True
        if self.superscript:
            run.font.superscript = True

    def render_blank_line(self, ast):
        pass

    def render_quote(self, ast):
        self.style_stack.append("Quote")
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_code_span(self, ast):
        self.current_paragraph.add_run(ast["children"]).style = self.doc.styles[
            "Macro Text Char"
        ]

    def render_code_block(self, ast):
        self.current_paragraph = self.doc.add_paragraph(style="Code")
        self.style_stack.append("Code")
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_fenced_code(self, ast):
        self.current_paragraph = self.doc.add_paragraph(style="Code")
        self.style_stack.append("Code")
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_image(self, ast):
        try:
            # Fetch image from the web.
            if urllib.parse.urlparse(ast["dest"]).scheme:
                response = requests.get(ast["dest"])
                if response.status_code != HTTP.OK:
                    raise ValueError(f"could not fetch image '{ast['dest']}'")
                if response.headers["Content-Type"] not in (
                    constants.PNG_MIMETYPE,
                    constants.JPEG_MIMETYPE,
                ):
                    raise ValueError(
                        f"cannot handle image '{ast['dest']}' with content type '{response.headers['Content-Type']}'"
                    )
                self.add_image(response.content, ast, 1.0)

            # Use image from the image library.
            elif ast["dest"] in get_imgs():
                img = get_imgs()[ast["dest"]]
                scale_factor = img["docx"]["scale_factor"]

                if img["content_type"] in (
                    constants.SVG_MIMETYPE,
                    constants.JSON_MIMETYPE,
                ):
                    # SVG image.
                    if img["content_type"] == constants.SVG_MIMETYPE:
                        # SVG in image library has already been checked for validity.
                        root = minixml.parse_content(img["data"])

                    # Vega-Lite plot.
                    else:
                        # JSON in image library has already been checked for validity.
                        vl_spec = json.loads(img["data"])
                        root = minixml.parse_content(
                            vl_convert.vegalite_to_svg(vl_spec)
                        )

                    # Set viewbox so that scaling behaves.
                    root["viewBox"] = f"0 0 {root['width']} {root['height']}"

                    # Scale width and height in SVG element.
                    rendering_factor = img["docx"]["png_rendering_factor"]
                    root["width"] = rendering_factor * float(root["width"])
                    root["height"] = rendering_factor * float(root["height"])

                    self.add_image(
                        vl_convert.svg_to_png(repr(root)),
                        ast,
                        scale_factor / rendering_factor,
                    )

                # JPEG or PNG.
                elif img["content_type"] in (
                    constants.PNG_MIMETYPE,
                    constants.JPEG_MIMETYPE,
                ):
                    self.add_image(
                        base64.standard_b64decode(img["data"]), ast, scale_factor
                    )
                else:
                    raise ValueError(f"cannot handle image {img['content_type']}")
            else:
                raise ValueError(f"no such image '{ast['dest']}'")

        except ValueError as error:
            self.current_paragraph = self.doc.add_paragraph(style="Code")
            self.current_paragraph.add_run(str(error))

    def render_emphasis(self, ast):
        self.italic = True
        for child in ast["children"]:
            self.render(child)
        self.italic = False

    def render_strong_emphasis(self, ast):
        self.bold = True
        for child in ast["children"]:
            self.render(child)
        self.bold = False

    def render_subscript(self, ast):
        self.subscript = True
        for child in ast["children"]:
            self.render(child)
        self.subscript = False

    def render_superscript(self, ast):
        self.superscript = True
        for child in ast["children"]:
            self.render(child)
        self.superscript = False

    def render_emdash(self, ast):
        self.current_paragraph.add_run(constants.EM_DASH)

    def render_line_break(self, ast):
        if ast.get("soft"):
            self.current_paragraph.add_run(" ")
        else:
            self.current_paragraph.add_run("\n")

    def render_thematic_break(self, ast):
        paragraph = self.doc.add_paragraph(constants.EM_DASH * 20)
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def render_link(self, ast):
        # This handles only raw text within a link; can't handle multiple children.
        raw_text = []
        for child in ast["children"]:
            if child["element"] == "raw_text":
                raw_text.append(child["children"])
        self.add_hyperlink(self.current_paragraph, ast["dest"], "".join(raw_text))

    def render_list(self, ast):
        data = dict(
            ordered=ast["ordered"],
            bullet=ast["bullet"],  # Currently not used.
            start=ast["start"],  # Currently not used.
            tight=ast["tight"],  # Currently not used.
            count=0,  # Currently not used.
            levels=len(self.list_stack) + 1,
        )
        self.list_stack.append(data)
        for child in ast["children"]:
            self.render(child)
        self.list_stack.pop()

    def render_list_item(self, ast):
        data = self.list_stack[-1]
        data["count"] += 1  # Currently not used.
        data["first_paragraph"] = True
        for child in ast["children"]:
            self.render(child)

    def render_indexed(self, ast):
        run = self.current_paragraph.add_run(ast["term"])
        if self.underline_indexed:
            run.font.underline = True
        if self.bold:
            run.font.bold = True
        if self.italic:
            run.font.italic = True

    def render_footnote_ref(self, ast):
        run = self.current_paragraph.add_run(f" {ast['number']}")
        run.font.superscript = True
        run.font.bold = True

    def render_footnote_def(self, ast):
        "The footnote definition in the element stream is not used; ignore."
        pass

    def render_reference(self, ast):
        self.current_paragraph.add_run(ast["name"]).font.bold = True
        self.current_paragraph.add_run(": ")
        reference = self.referenced[ast["name"]]
        self.current_paragraph.add_run(reference["title"]).font.italic = True

    def render_internal_link(self, ast):
        run = self.current_paragraph.add_run(ast["section"])
        run.font.color.rgb = docx.shared.RGBColor(225, 0, 0)
